import os
import asyncio
import subprocess
from dotenv import load_dotenv
from typing import Optional, Union, Iterable
from pyzerox import zerox  # 从 pyzerox 导入 zerox 函数
import markdown2
from docx import Document
from docx.shared import Pt

# 从 .env 文件加载 API 密钥
load_dotenv()  # 加载当前目录下的 .env 文件
os.environ["OPENAI_API_KEY"] = os.getenv(
    "OPENAI_API_KEY"
)  # 从 .env 文件获取 OPENAI_API_KEY


# 将 Markdown 转换为 DOCX 的函数
def markdown_to_docx(markdown_content: str, output_path: str):
    try:
        # 1. 将 Markdown 转换为 HTML
        html_content = markdown2.markdown(markdown_content)
        print(
            f"Markdown to HTML conversion successful. HTML content preview: {html_content[:500]}..."
        )

        # 2. 创建一个新的 DOCX 文件
        doc = Document()

        # 3. 使用 BeautifulSoup 解析 HTML
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # 4. 遍历 HTML 元素并转换为 DOCX
        for element in soup.find_all(["h1", "h2", "h3", "p", "strong", "a"]):
            if element.name == "h1":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "p":
                p = doc.add_paragraph(element.get_text())
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(12)
            elif element.name == "strong":
                p = doc.add_paragraph()
                p.add_run(element.get_text()).bold = True
            elif element.name == "a":
                p = doc.add_paragraph()
                p.add_run(element.get_text()).italic = True
                p.add_run(f"({element['href']})").italic = True
            elif element.name == "table":
                table = doc.add_table(rows=0, cols=0)
                table.style = 'Table Grid'
                for row in element.find_all("tr"):
                    cells = row.find_all(["td", "th"])
                    if not table.columns:
                        table.add_column(width=Pt(100))
                    row_cells = table.add_row().cells
                    for idx, cell in enumerate(cells):
                        if idx >= len(row_cells):
                            table.add_column(width=Pt(100))
                            row_cells = table.add_row().cells
                        row_cells[idx].text = cell.get_text()

        # 5. 保存 DOCX 文件
        doc.save(output_path)
        print(f"Document saved to {output_path}")
    except Exception as e:
        print(f"Error occurred while processing the document: {e}")


# 处理文件的函数
async def process_model():
    # 设置输入输出目录
    input_dir = "inputfile"
    output_dir = "outputfile"

    # 获取输入文件夹中的所有 PDF 文件
    pdf_files = [f for f in os.listdir(input_dir) if f.endswith(".pdf")]

    # 遍历每个 PDF 文件并进行处理
    for pdf_file in pdf_files:
        input_file_path = os.path.join(input_dir, pdf_file)
        output_file_path = os.path.join(
            output_dir, f"{os.path.splitext(pdf_file)[0]}.docx"
        )

        print(f"Processing file: {input_file_path}")

        try:
            # 1. 使用 Zerox 提取 PDF 内容
            result = await zerox(file_path=input_file_path, model="gpt-4o-mini")

            # 2. 获取处理后的 Markdown 内容（整合所有页面）
            markdown_content = "\n".join([page.content for page in result.pages])

            # 另外保存为 .md 文件
            md_output_file_path = os.path.join(
                output_dir, f"{os.path.splitext(pdf_file)[0]}.md"
            )
            with open(md_output_file_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)

            # 3. 将 Markdown 转换为 DOCX 并保存
            markdown_to_docx(markdown_content, output_file_path)

        except Exception as e:
            print(f"Error processing {pdf_file}: {e}")


# 主函数
def main():
    asyncio.run(process_model())


if __name__ == "__main__":
    main()